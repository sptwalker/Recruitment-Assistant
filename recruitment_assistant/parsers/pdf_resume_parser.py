import re
from dataclasses import dataclass, field
from decimal import Decimal
from pathlib import Path

from docx import Document
from pypdf import PdfReader

PHONE_RE = re.compile(r"(?<!\d)(?:1[3-9]\d{9})(?!\d)")
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
DEGREES = ["博士", "硕士", "研究生", "本科", "大专", "专科", "高中", "中专"]
SECTION_KEYWORDS = [
    "个人信息", "基本信息", "个人优势", "求职意向", "工作经历", "工作经验", "实习经历", "项目经历", "项目经验",
    "教育经历", "教育背景", "自我评价", "技能", "专业技能", "技能证书", "资质证书", "证书", "项目亮点",
]
NOISE_KEYWORDS = [
    "在线", "沟通", "聊天", "未读", "未联系", "已联系", "打招呼", "要附件简历", "已向对方要附件简历",
    "查看简历", "查看附件简历", "设置备注", "不合适", "待识别",
    "快速处理", "新招呼", "全部职位", "筛选", "批量", "智联", "简历", "附件", "下载",
]
JOB_HINTS = [
    "工程师", "开发", "经理", "主管", "专员", "助理", "运营", "产品", "设计", "销售", "财务", "人事", "行政", "顾问", "总监",
    "算法", "测试", "前端", "后端", "架构", "实施", "运维", "会计", "出纳", "法务", "教师", "司机", "客服", "Unity", "Java", "Python",
    "AI", "大模型", "人工智能", "UI", "交互", "美工", "平面", "视觉", "数据", "机器学习", "深度学习",
]
KNOWN_CITIES = {"广州", "深圳", "上海", "北京", "杭州", "南京", "苏州", "成都", "重庆", "武汉", "长沙", "西安", "盐城"}
STATUS_WORDS = ["离职", "在职", "全职", "兼职", "正在找工作", "求职状态", "目前状态", "随时到岗"]
CITY_SUFFIX_RE = re.compile(r"([\u4e00-\u9fa5]{2,12}(?:省|市|区|县))")
DATE_LINE_RE = re.compile(r"(?:19|20)\d{2}[./年/-]\s*\d{1,2}|至今")



@dataclass
class ParsedResume:
    source_file: str
    text: str
    name: str | None = None
    job_title: str | None = None
    phone: str | None = None
    email: str | None = None
    current_city: str | None = None
    highest_degree: str | None = None
    years_of_experience: Decimal | None = None
    current_company: str | None = None
    current_position: str | None = None
    expected_position: str | None = None
    skills: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "source_file": self.source_file,
            "name": self.name,
            "job_title": self.job_title,
            "phone": self.phone,
            "email": self.email,
            "current_city": self.current_city,
            "highest_degree": self.highest_degree,
            "years_of_experience": str(self.years_of_experience) if self.years_of_experience else None,
            "current_company": self.current_company,
            "current_position": self.current_position,
            "expected_position": self.expected_position,
            "skills": self.skills,
            "text_preview": self.text[:1000],
        }


def normalize_text(text: str) -> str:
    lines = []
    for line in text.replace("\r", "\n").split("\n"):
        cleaned = re.sub(r"\s+", " ", line).strip()
        if cleaned:
            lines.append(cleaned)
    return "\n".join(lines)


def extract_pdf_text(file_path: str | Path) -> str:
    reader = PdfReader(str(file_path))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return normalize_text("\n".join(parts))


def extract_docx_text(file_path: str | Path) -> str:
    document = Document(str(file_path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return normalize_text("\n".join(parts))


def extract_doc_text(file_path: str | Path) -> str:
    data = Path(file_path).read_bytes()
    candidates = []
    for encoding in ("utf-16le", "gb18030", "utf-8"):
        try:
            text = data.decode(encoding, errors="ignore")
        except Exception:
            continue
        cleaned = "".join(
            ch if (ch == "\n" or ch == "\t" or (ch.isprintable() and not 0xD800 <= ord(ch) <= 0xDFFF)) else "\n"
            for ch in text
        )
        cleaned = normalize_text(cleaned)
        score = sum(1 for ch in cleaned if "\u4e00" <= ch <= "\u9fff")
        if score:
            candidates.append((score, cleaned))
    return max(candidates, key=lambda item: item[0])[1] if candidates else ""


def parse_resume_file(file_path: str | Path, candidate_signature: str | None = None) -> ParsedResume:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = extract_pdf_text(path)
    elif suffix == ".docx":
        text = extract_docx_text(path)
    elif suffix == ".doc":
        text = extract_doc_text(path)
    else:
        raise ValueError(f"暂不支持解析 {suffix} 文件")
    return parse_resume_text(path, text, candidate_signature=candidate_signature)


def parse_resume_pdf(file_path: str | Path) -> ParsedResume:
    return parse_resume_file(file_path)


def parse_resume_text(path: Path, text: str, candidate_signature: str | None = None) -> ParsedResume:
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    parsed = ParsedResume(source_file=str(path), text=text)
    parsed.phone = find_phone(text)
    parsed.email = find_email(text)
    parsed.name = find_name(lines, parsed.phone, parsed.email)
    parsed.highest_degree = find_highest_degree(text)
    parsed.years_of_experience = find_years_of_experience(text)
    parsed.current_city = find_city(lines)
    parsed.expected_position = find_expected_position(lines)
    parsed.current_company, parsed.current_position = find_current_work(lines)
    parsed.skills = find_skills(text)

    name_from_file, job_from_file = infer_name_and_job_from_filename(path.name)
    name_from_sig, _ = clean_candidate_signature(candidate_signature or "")
    parsed.name = first_valid_name(find_name_near_contact(lines, parsed.phone, parsed.email), parsed.name, name_from_file, name_from_sig)
    parsed.current_city = parsed.current_city or find_city(lines)
    parsed.expected_position = parsed.expected_position or find_expected_position_by_label(lines)
    fallback_company, fallback_position = find_recent_work(lines)
    if not parsed.current_company or not is_company_name(parsed.current_company):
        parsed.current_company = fallback_company
    if not parsed.current_position or not is_valid_job_title(parsed.current_position):
        parsed.current_position = fallback_position
    parsed.job_title = first_non_empty(parsed.expected_position, parsed.current_position, job_from_file)
    return parsed



def find_job_title_from_header(lines: list[str], name: str | None = None) -> str | None:
    skip_prefixes = ("研究方向", "专业方向", "主修", "毕业", "学历", "年龄", "电话", "邮箱", "现所在地", "联系信息")
    skip_contains = ("教育经历", "工作经历", "项目经历", "专业", "大学", "学院", "本科", "硕士", "博士", "大专", "毕业时间")
    invalid_chars = "。；;，,、"
    for line in lines[:18]:
        line = line.strip(" -—｜|")
        if not line or (name and line == name):
            continue
        if line.startswith(skip_prefixes) or any(token in line for token in skip_contains):
            continue
        if len(line) <= 40 and not any(char in line for char in invalid_chars) and any(hint.lower() in line.lower() for hint in JOB_HINTS):
            return line
        for part in split_resume_tokens(line):
            if part == name or first_valid_name(part) == part:
                continue
            if len(part) <= 30 and not any(char in part for char in invalid_chars) and any(hint.lower() in part.lower() for hint in JOB_HINTS):
                return part
    return None


def clean_candidate_signature(signature: str) -> tuple[str | None, str | None]:
    if not signature:
        return None, None
    text = re.sub(r"\s+", " ", signature).strip()
    text = re.sub(r"^\s*\d+\s+", "", text)
    text = re.sub(r"pos:\d+:\d+:\d+:\d+", " ", text)
    parts = split_resume_tokens(text)
    parts = [part for part in parts if part and not is_noise_token(part)]
    name = first_valid_name(*parts[:6])
    job = None
    for part in parts:
        if part == name or re.search(r"\d+岁|\d+年|本科|硕士|博士|大专|深圳|广州|上海|北京", part):
            continue
        if any(hint in part for hint in JOB_HINTS) and 2 <= len(part) <= 40:
            job = part
            break
    return name, job


def infer_name_and_job_from_filename(filename: str) -> tuple[str | None, str | None]:
    stem = Path(filename).stem
    stem = re.sub(r"^zhilian_\d{8}_\d{6}_", "", stem)
    stem = re.sub(r"_[0-9a-f]{8,}$", "", stem, flags=re.I)
    parts = split_resume_tokens(stem)
    parts = [part for part in parts if part and not is_noise_token(part)]
    name = first_valid_name(*parts[:4])
    job = None
    for part in parts:
        if part == name or re.fullmatch(r"\d+岁", part) or part in {"深圳", "广州", "上海", "北京"}:
            continue
        cleaned = clean_job_title(part)
        if cleaned:
            job = cleaned
            break
    return name, job


def split_resume_tokens(text: str) -> list[str]:
    text = re.sub(r"[()（）\[\]【】]", " ", text)
    return [item.strip(" _-—|｜/\\,，;；:") for item in re.split(r"[_\-|｜/\\,，;；\n\t ]+", text) if item.strip()]


def is_noise_token(value: str) -> bool:
    if any(keyword in value for keyword in NOISE_KEYWORDS):
        return True
    if re.fullmatch(r"[0-9a-f]{12,}", value, re.I):
        return True
    if re.fullmatch(r"\d{4,}", value):
        return True
    return False


def first_non_empty(*values: str | None) -> str | None:
    for value in values:
        if value and str(value).strip():
            return str(value).strip()
    return None


def first_valid_name(*values: str | None) -> str | None:
    for value in values:
        if not value:
            continue
        text = str(value).strip()
        text = re.sub(r"^[A-Za-z](?=[\u4e00-\u9fa5·]{2,8}$)", "", text)
        labeled = re.search(r"(?:姓\s*名|姓名)\s*[:：]?\s*([\u4e00-\u9fa5·]{2,8})", text)
        if labeled:
            text = labeled.group(1)
        cleaned = re.sub(r"[|｜_/\\,，;；:].*", "", text).strip()
        cleaned = re.sub(r"(?:年龄|学\s*历|学历|性\s*别|性别|专\s*业|专业).*", "", cleaned).strip()
        if cleaned in KNOWN_CITIES:
            continue
        if any(keyword in cleaned for keyword in NOISE_KEYWORDS + JOB_HINTS + SECTION_KEYWORDS + STATUS_WORDS):
            continue
        if re.fullmatch(r"[\u4e00-\u9fa5·]{2,8}", cleaned):
            return cleaned
    return None




def find_name_near_contact(lines: list[str], phone: str | None, email: str | None) -> str | None:
    for index, line in enumerate(lines):
        if not ((phone and phone in line) or (email and email in line)):
            continue
        same_line = re.sub(PHONE_RE, " ", line)
        same_line = re.sub(EMAIL_RE, " ", same_line)
        same_line = re.sub(r"\b\d{1,2}\s*岁\b", " ", same_line)
        for part in split_resume_tokens(same_line):
            name = first_valid_name(part)
            if name:
                return name
        for offset in (-3, -2, -1, 1, 2, 3):
            pos = index + offset
            if 0 <= pos < len(lines):
                name = first_valid_name(lines[pos])
                if name:
                    return name
    for line in reversed(lines[-10:]):
        name = first_valid_name(line)
        if name:
            return name
    return None


def find_labeled_value(lines: list[str], labels: list[str], max_len: int = 40) -> str | None:
    label_pattern = "|".join(re.escape(label) for label in labels)
    for index, line in enumerate(lines):
        if not any(label in line for label in labels):
            continue
        matches = list(re.finditer(rf"(?:{label_pattern})\s*[:：]?\s*([^:：\n]+?)(?=\s*(?:工作性质|期望薪资|期望职位|期望职业|目前状态|到岗时间|目标地点|$))", line, flags=re.I))
        for match in matches:
            raw_value = match.group(1).strip()
            value = clean_job_title(raw_value) or raw_value.strip(" -—｜|:：")
            if value and len(value) <= max_len and value not in {"：", ":"} and not any(label == value for label in labels):
                return value
        if index + 1 < len(lines):
            value = clean_job_title(lines[index + 1]) or lines[index + 1].strip(" -—｜|:：")
            if value and len(value) <= max_len and value not in {"：", ":"} and not any(label == value for label in labels):
                return value
    return None


def extract_position_from_line(line: str) -> str | None:
    matches = re.findall(
        r"(?:期望职位|期望职业|应聘职位|目标职位)\s*[:：]?\s*(.+?)(?=\s*(?:工作性质|期望薪资|目前状态|到岗时间|目标地点|期望行业|$))",
        line,
    )
    for value in matches:
        cleaned = clean_job_title(value.strip()) or value.strip(" -—｜|:：")
        if cleaned and cleaned not in {"期望职位", "期望职业"}:
            return cleaned
    return None





def find_expected_position_by_label(lines: list[str]) -> str | None:
    return find_labeled_value(lines, ["期望职位", "期望职业", "应聘职位", "目标职位", "求职岗位", "岗位名称"], max_len=60)


def find_recent_work(lines: list[str]) -> tuple[str | None, str | None]:
    for index, line in enumerate(lines):
        if "公司名称" in line or "岗位名称" in line:
            company_match = re.search(r"公司名称[:：]?\s*([^\s].*?)(?:\s+岗位名称|$)", line)
            position_match = re.search(r"岗位名称[:：]?\s*(.+)$", line)
            company = company_match.group(1).strip() if company_match else None
            position = clean_job_title(position_match.group(1)) if position_match else None
            if company or position:
                return company, position
    in_work_section = False
    for index, line in enumerate(lines):
        if any(section in line for section in ["工作经历", "工作经验", "实习经历"]):
            in_work_section = True
            continue
        if in_work_section and any(section in line for section in ["项目经历", "项目经验", "教育经历", "教育背景", "自我评价", "技能", "证书"]):
            break
        if not in_work_section:
            continue
        if DATE_LINE_RE.search(line):
            company, position = parse_work_line(line)
            if company or position:
                return company, position
            next_lines = [item for item in lines[index + 1:index + 4] if item.strip()]
            company = next((item.strip() for item in next_lines if is_company_name(item)), None)
            position = next((clean_job_title(item) for item in next_lines if clean_job_title(item)), None)
            if company or position:
                return company, position
    return None, None





def is_company_name(value: str | None) -> bool:
    text = str(value or "").strip()
    if not text or len(text) > 80:
        return False
    if any(token in text for token in ["大学", "学院", "学校", "本科", "硕士", "博士", "主修课程"]):
        return False
    return any(token in text for token in ["有限公司", "公司", "集团", "科技", "网络", "通信", "工作室", "淘宝", "中心", "银行"])


def find_phone(text: str) -> str | None:
    match = PHONE_RE.search(text)
    return match.group(0) if match else None








def find_email(text: str) -> str | None:
    match = EMAIL_RE.search(text)
    return match.group(0) if match else None


def find_name(lines: list[str], phone: str | None, email: str | None) -> str | None:
    for line in lines[:30]:
        if phone and phone in line:
            continue
        if email and email in line:
            continue
        if any(keyword in line for keyword in SECTION_KEYWORDS + NOISE_KEYWORDS + JOB_HINTS + STATUS_WORDS):
            continue
        if re.search(r"年龄|性别|工作年限|求职|期望|工作地区|月薪|行业|http|www|@", line, re.I):
            continue
        for part in split_resume_tokens(line):
            name = first_valid_name(part)
            if name:
                return name
    return None


def find_highest_degree(text: str) -> str | None:
    for degree in DEGREES:
        if degree in text:
            if degree == "研究生":
                return "硕士"
            if degree == "专科":
                return "大专"
            return degree
    return None


def find_years_of_experience(text: str) -> Decimal | None:
    for pattern in [r"(\d+(?:\.\d+)?)\s*年(?:以上)?工作经验", r"工作年限[:：]?\s*(\d+(?:\.\d+)?)\s*年", r"工作经验[:：]?\s*(\d+(?:\.\d+)?)\s*年", r"(\d+(?:\.\d+)?)\s*年经验"]:
        match = re.search(pattern, text)
        if match:
            years = Decimal(match.group(1))
            return years if Decimal("0") <= years <= Decimal("80") else None
    return None


def find_city(lines: list[str]) -> str | None:
    city_noise = {"求职状态", "工作性质", "期望月薪", "期望职业", "期望行业", "教育经历", "工作经历", "项目经历", "工作地区"}
    for index, line in enumerate(lines[:80]):
        if any(label in line for label in ["工作地区", "目标地点", "当前城市", "居住地", "现居", "所在地"]):
            same_line = re.sub(r".*?(工作地区|目标地点|当前城市|居住地|现居|所在地)[:： ]*", "", line).strip()
            for city in KNOWN_CITIES:
                if city in same_line:
                    return city
            if index + 1 < len(lines) and lines[index + 1].strip() in KNOWN_CITIES:
                return lines[index + 1].strip()
    for line in lines[:80]:
        match = re.search(r"(?:籍\s*贯|户籍|现居|所在地|当前城市|居住地|城市)[:： ]*([\u4e00-\u9fa5]{2,12})", line)
        if match:
            city = match.group(1).strip()
            if city not in city_noise and city not in {"大学"}:
                return city
    for line in lines[:80]:
        if any(token in line for token in city_noise):
            continue
        for city in KNOWN_CITIES:
            if re.search(rf"(?:^|[\s|｜/]){city}(?:$|[\s|｜/])", line):
                return city
        city_match = CITY_SUFFIX_RE.search(line)
        if city_match:
            city = city_match.group(1)
            if city not in {"大学"}:
                return city
        if line in KNOWN_CITIES:
            return line
    return None



def clean_job_title(value: str | None) -> str | None:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" -—｜|:：")
    if not text:
        return None
    company_noise = ["有限公司", "分公司", "集团", "科技", "公司", "企业", "中心", "事业部", "工作室", "系统集成"]
    section_noise = ["工作经历", "工作经验", "项目经历", "项目经验", "教育经历", "教育背景", "实习经历", "培训经历", "校园经历", "主修课程", "专业技能", "工作描述", "项目描述"]
    if any(token in text for token in section_noise):
        return None
    text = re.sub(r"^(期望职位|期望职业|求职意向|应聘职位|求职岗位|应聘岗位|目标职位|目标岗位|职位|岗位)[:： ]*", "", text).strip(" -—｜|")
    text = re.split(r"\s+(?:工作性质|期望薪资|目前状态|到岗时间|目标地点|期望行业)[:：]?", text)[0].strip(" -—｜|")
    text = re.split(r"电话|手机|性别|姓名|男|女|\d{2,}|岁|经验|本科|专科|硕士|博士|学历|在线|沟通|附件|简历", text)[0].strip(" -—｜|")
    text = re.sub(r"^(?:19|20)\d{2}(?:[./-]\s*\d{1,2}|年\s*\d{1,2}月?)?\s*(?:[-—–至到]|年)?\s*(?:(?:19|20)\d{2}(?:[./-]\s*\d{1,2}|年\s*\d{1,2}月?)?|至今)?\s*", "", text).strip(" -—｜|")
    company_match = re.search(r"[\u4e00-\u9fa5A-Za-z0-9（）()·\-]{2,40}(?:有限公司|公司|集团|科技有限公司|网络信息科技有限公司|通信科技有限公司|工作室|淘宝服装店)", text)
    if company_match:
        text = text.replace(company_match.group(0), " ", 1).strip(" -—｜|")
    parts = [part.strip(" -—｜|/\\,，;；:：()（）[]【】") for part in re.split(r"[·•|｜/\\,，;；\n\r\t]+", text)]
    candidates = [part for part in parts if part]
    candidates.append(text)
    for part in reversed(candidates):
        if is_valid_job_title(part):
            return part
    return text if is_valid_job_title(text) and not any(token in text for token in company_noise + section_noise) else None


def is_valid_job_title(value: str | None) -> bool:
    text = str(value or "").strip()
    if not (2 <= len(text) <= 40):
        return False
    invalid_tokens = [
        "主修课程", "工作描述", "项目描述", "负责", "参与", "熟练", "具备", "提升", "增强", "数据清洗", "课程",
        "大学", "学院", "本科", "硕士", "博士", "研究方向", "专业排名",
    ]
    if any(token in text for token in invalid_tokens):
        return False
    if "系统" in text and not any(token in text for token in ["工程师", "开发", "运维", "架构", "经理", "主管"]):
        return False
    return any(hint.lower() in text.lower() for hint in JOB_HINTS)




def find_expected_position(lines: list[str]) -> str | None:
    for index, line in enumerate(lines):
        if any(label in line for label in ["期望职位", "期望职业", "求职意向", "应聘职位", "目标职位"]):
            extracted = extract_position_from_line(line)
            if extracted:
                return extracted
            value = re.sub(r".*?(期望职位|期望职业|求职意向|应聘职位|目标职位)[:： ]*", "", line).strip()
            cleaned = clean_job_title(value)
            if cleaned:
                return cleaned
            if value in {"", "：", ":"} and index + 1 < len(lines):
                return clean_job_title(lines[index + 1])
    return None





def find_current_work(lines: list[str]) -> tuple[str | None, str | None]:
    for index, line in enumerate(lines):
        if "工作经历" not in line and "工作经验" not in line and "实习经历" not in line:
            continue
        for offset, work_line in enumerate(lines[index + 1:index + 12], start=1):
            if any(section in work_line for section in ["项目经历", "项目经验", "教育经历", "教育背景", "自我评价", "技能", "证书"]):
                break
            if "公司名称" in work_line or "岗位名称" in work_line:
                company_match = re.search(r"公司名称[:：]?\s*([^\s].*?)(?:\s+岗位名称|$)", work_line)
                position_match = re.search(r"岗位名称[:：]?\s*(.+)$", work_line)
                company = company_match.group(1).strip() if company_match else None
                position = clean_job_title(position_match.group(1)) if position_match else None
                if company or position:
                    return company, position
            if DATE_LINE_RE.search(work_line):
                company, position = parse_work_line(work_line)
                if company or position:
                    return company, position
                next_lines = [item for item in lines[index + offset + 1:index + offset + 4] if item.strip()]
                company = next((item.strip() for item in next_lines if is_company_name(item)), None)
                position = next((clean_job_title(item) for item in next_lines if clean_job_title(item)), None)
                if company or position:
                    return company, position
            if is_company_name(work_line):
                position = clean_job_title(lines[index + offset + 1]) if index + offset + 1 < len(lines) else None
                return work_line.strip(), position
    return None, None


def parse_work_line(line: str) -> tuple[str | None, str | None]:
    cleaned = re.sub(r"^\s*(?:19|20)\d{2}(?:[./-]\s*\d{1,2}|年\s*\d{1,2}月?)?\s*(?:[-—–至到]|年)?\s*(?:(?:19|20)\d{2}(?:[./-]\s*\d{1,2}|年\s*\d{1,2}月?)?|至今)?\s*", "", line).strip(" -—｜|")
    if not cleaned:
        return None, None
    company_match = re.search(r"([\u4e00-\u9fa5A-Za-z0-9（）()·\-]{2,40}(?:有限公司|公司|集团|科技有限公司|网络信息科技有限公司|通信科技有限公司|学院|银行|工作室|淘宝服装店))", cleaned)
    company = company_match.group(1).strip() if company_match else (cleaned if is_company_name(cleaned) else None)
    rest = cleaned.replace(company, " ", 1).strip(" -—｜|") if company else cleaned
    position = clean_job_title(rest)
    return company, position



def find_skills(text: str) -> list[str]:
    common_skills = [
        "Python", "Java", "C++", "C#", "JavaScript", "TypeScript", "React", "Vue", "SQL", "MySQL", "PostgreSQL", "Oracle", "Redis", "Docker", "Kubernetes", "Linux", "Excel",
        "Photoshop", "Illustrator", "Sketch", "Axure", "Figma", "CAD", "SolidWorks", "PLC", "PMP", "销售", "运营", "招聘", "人事", "财务",
        "机器学习", "深度学习", "计算机视觉", "自然语言处理", "YOLO", "Transformer", "PyTorch", "TensorFlow", "OpenCV", "LangChain", "LangGraph", "RAG", "大模型", "Agent",
        "MongoDB", "Milvus", "Chroma", "HTML", "CSS", "DREAMWEAVER", "FLASH", "Coreldraw", "MasterGo", "PyQt", "Gradio",
    ]
    lower_text = text.lower()
    return [skill for skill in common_skills if skill.lower() in lower_text]
