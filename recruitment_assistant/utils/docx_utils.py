"""DOCX 文本提取工具。

部分简历模板把全部正文塞进 Word 文本框（<w:txbxContent>），python-docx 的
doc.paragraphs / doc.tables 只遍历 body 顶层，不下钻文本框，会返回空文本。
本模块提供 docx_xml_text_fallback：直接读 zip 内 word/document.xml + 所有
header*.xml / footer*.xml，正则抓所有 <w:t> 拼接 — 兜底覆盖文本框 / SDT /
drawing 等结构。
"""
from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document

_WT_PATTERN = re.compile(r"<w:t(?:\s[^>]*)?>([^<]*)</w:t>")


def docx_xml_text_fallback(file_path: str | Path) -> str:
    """直接从 docx 的 zip 内部 XML 抓所有 <w:t> 文本，覆盖文本框等 paragraphs 不可达的位置。"""
    path = Path(file_path)
    try:
        with zipfile.ZipFile(str(path), "r") as zf:
            names = zf.namelist()
            xml_parts: list[str] = []
            if "word/document.xml" in names:
                xml_parts.append(zf.read("word/document.xml").decode("utf-8", errors="ignore"))
            for n in names:
                if (n.startswith("word/header") or n.startswith("word/footer")) and n.endswith(".xml"):
                    try:
                        xml_parts.append(zf.read(n).decode("utf-8", errors="ignore"))
                    except Exception:
                        continue
    except Exception:
        return ""
    if not xml_parts:
        return ""
    combined = "\n".join(xml_parts)
    texts = _WT_PATTERN.findall(combined)
    return "\n".join(t for t in (s.strip() for s in texts) if t)


def extract_docx_text(file_path: str | Path) -> str:
    try:
        document = Document(str(file_path))
    except Exception:
        return docx_xml_text_fallback(file_path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    result = "\n".join(paragraphs)
    if len(result.strip()) < 50:
        fallback = docx_xml_text_fallback(file_path)
        if len(fallback.strip()) > len(result.strip()):
            return fallback
    return result
